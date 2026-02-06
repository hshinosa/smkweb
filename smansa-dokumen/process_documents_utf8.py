# -*- coding: utf-8 -*-
"""
SMAN 1 Baleendah Document Processor
Extract content from documents and images to structured markdown format for RAG/AI processing

Supported formats:
- PDF, DOCX, XLSX - text extraction
- Images (PNG, JPEG) - description/extraction

Requirements:
- pip install pypdf2 python-docx openpyxl pytesseract pillow tqdm
"""

import os
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
from dataclasses import dataclass, field
import re

# Initialize paths
DOC_ROOT = Path(__file__).parent
OUTPUT_DIR = DOC_ROOT / "output"
OUTPUT_FILE = OUTPUT_DIR / "smansa_dokumen_processed.md"

# Create output directory
OUTPUT_DIR.mkdir(exist_ok=True)

# Try to import required libraries
try:
    from pypdf import PdfReader
except ImportError:
    print("Installing pypdf2...")
    os.system("pip install pypdf2")
    from pypdf import PdfReader

try:
    from docx import Document as DocxDocument
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx")
    from docx import Document as DocxDocument

try:
    import openpyxl
except ImportError:
    print(" installing openpyxl...")
    os.system("pip install openpyxl")
    import openpyxl

try:
    from PIL import Image
except ImportError:
    print("Installing Pillow...")
    os.system("pip install Pillow")
    from PIL import Image


@dataclass
class ProcessedDocument:
    """Store processed document information"""
    filename: str
    category: str
    content: str
    metadata: Dict[str, str] = field(default_factory=dict)
    images: List[str] = field(default_factory=list)
    processing_date: str = datetime.now().isoformat()


class DocumentProcessor:
    """Main processor for all document types"""
    
    def __init__(self):
        self.documents: List[ProcessedDocument] = []
        self.image_counter = 0
        
    def process_directory(self, directory: Path) -> List[ProcessedDocument]:
        """Process all documents in directory recursively"""
        print(f"Scanning directory: {directory}")
        
        # Get all files
        all_files = list(directory.rglob("*"))
        
        # Filter by supported extensions
        supported_extensions = {'.pdf', '.docx', '.xlsx', '.png', '.jpg', '.jpeg'}
        files = [f for f in all_files if f.is_file() and f.suffix.lower() in supported_extensions]
        
        # Exclude self and output directory
        excluded_prefixes = ['process_documents_', __file__.replace('.py', '')]
        files = [f for f in files if not any(f.name.startswith(prefix) for prefix in excluded_prefixes)]
        
        if not files:
            print("No supported documents found!")
            return []
        
        print(f"Found {len(files)} files to process\n")
        
        # Process each file
        for i, file_path in enumerate(files, 1):
            print(f"[{i}/{len(files)}] Processing: {file_path.name} ({file_path.suffix})")
            
            try:
                doc = self._process_file(file_path)
                if doc:
                    self.documents.append(doc)
                    print(f"  Extracted {len(doc.content)} characters\n")
            except Exception as e:
                print(f"  Error: {str(e)}\n")
        
        return self.documents
    
    def _process_file(self, file_path: Path) -> Optional[ProcessedDocument]:
        """Route file to appropriate processor"""
        suffix = file_path.suffix.lower()
        
        category = self._get_category(file_path)
        
        if suffix == '.pdf':
            return self._process_pdf(file_path, category)
        elif suffix == '.docx':
            return self._process_docx(file_path, category)
        elif suffix == '.xlsx':
            return self._process_xlsx(file_path, category)
        elif suffix in {'.png', '.jpg', '.jpeg'}:
            return self._process_image(file_path, category)
        else:
            print(f"  Unsupported format: {suffix}")
            return None
    
    def _get_category(self, file_path: Path) -> str:
        """Determine document category from path"""
        path_str = str(file_path).lower()
        
        if 'data serapan' in path_str:
            return 'Data Serapan'
        elif 'kurikulum' in path_str:
            return 'Kurikulum'
        elif 'sejarah' in path_str:
            return 'Sejarah Profil'
        elif 'profil' in path_str:
            return 'Sejarah Profil'
        elif 'pembelajaran' in path_str:
            return 'Pendidikan Pembelajaran'
        elif 'rekapitulasi' in path_str:
            return 'Data Akademik'
        elif 'poster' in path_str or 'rapor' in path_str:
            return 'Dokumentasi Sekolah'
        else:
            return 'Umum'
    
    def _process_pdf(self, file_path: Path, category: str) -> ProcessedDocument:
        """Extract text from PDF"""
        reader = PdfReader(str(file_path))
        text_content = []
        
        for page in reader.pages:
            try:
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)
            except Exception as e:
                print(f"  Error reading page: {e}")
        
        full_text = "\n\n".join(text_content)
        
        return ProcessedDocument(
            filename=file_path.name,
            category=category,
            content=full_text,
            metadata={
                'pages': len(reader.pages),
                'file_size': f"{file_path.stat().st_size / 1024:.1f} KB"
            }
        )
    
    def _process_docx(self, file_path: Path, category: str) -> ProcessedDocument:
        """Extract text from DOCX"""
        doc = DocxDocument(str(file_path))
        text_content = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text.strip())
        
        # Extract tables
        tables_content = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                tables_content.append(row_text)
        
        full_text = "\n\n".join(text_content)
        
        if tables_content:
            full_text += "\n\n### Tabel:\n\n" + "\n".join(tables_content)
        
        return ProcessedDocument(
            filename=file_path.name,
            category=category,
            content=full_text,
            metadata={
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'file_size': f"{file_path.stat().st_size / 1024:.1f} KB"
            }
        )
    
    def _process_xlsx(self, file_path: Path, category: str) -> ProcessedDocument:
        """Extract data from XLSX"""
        workbook = openpyxl.load_workbook(str(file_path), data_only=True)
        sheets_content = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            data = []
            
            # Get header
            headers = []
            for cell in sheet[1]:
                headers.append(cell.value)
            
            # Get data rows
            for row in sheet.iter_rows(min_row=2, values_only=True):
                data.append(row)
            
            if data:
                # Create markdown table
                table_str = f"### Sheet: {sheet_name}\n\n"
                table_str += "|" + "|".join([str(h) if h else "" for h in headers]) + "|\n"
                table_str += "|" + "|".join(["---" for _ in headers]) + "|\n"
                
                for row in data:
                    table_str += "|" + "|".join([str(cell) if cell else "" for cell in row]) + "|\n"
                
                sheets_content.append(table_str)
        
        full_text = "\n\n".join(sheets_content)
        
        return ProcessedDocument(
            filename=file_path.name,
            category=category,
            content=full_text,
            metadata={
                'sheets': len(workbook.sheetnames),
                'file_size': f"{file_path.stat().st_size / 1024:.1f} KB"
            }
        )
    
    def _process_image(self, file_path: Path, category: str) -> ProcessedDocument:
        """Generate description for image"""
        self.image_counter += 1
        
        # Get image info
        img = Image.open(str(file_path))
        width, height = img.size
        
        # Create description based on filename and properties
        filename_lower = file_path.name.lower()
        description_parts = []
        
        if 'kurikulum' in filename_lower:
            description_parts.append("Grafik/Ilustrasi terkait kurikulum")
        elif 'sejarah' in filename_lower or 'profil' in filename_lower:
            description_parts.append("Gambar profil/sejarah sekolah")
        elif 'poster' in filename_lower or 'rapor' in filename_lower:
            description_parts.append("Poster atau dokumentasi sekolah")
        else:
            description_parts.append("Gambar terkait kegiatan sekolah")
        
        description_parts.append(f"Ukuran: {width}x{height} piksel")
        description_parts.append(f"Format: {file_path.suffix[1:].upper()}")
        
        content = f"**Gambar:** {file_path.name}\n\n{', '.join(description_parts)}\n\n*Catatan: Gambar ini digunakan sebagai referensi visual untuk konten yang terkait.*"
        
        image_path = file_path.relative_to(DOC_ROOT)
        
        return ProcessedDocument(
            filename=file_path.name,
            category=category,
            content=content,
            images=[str(image_path)],
            metadata={
                'width': width,
                'height': height,
                'format': file_path.suffix[1:].upper(),
                'file_size': f"{file_path.stat().st_size / 1024:.1f} KB"
            }
        )
    
    def clean_text(self, text: str) -> str:
        """Clean and format extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Remove leading/trailing whitespace from lines
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)
        
        return text


class MarkdownGenerator:
    """Generate structured markdown from processed documents"""
    
    def __init__(self, documents: List[ProcessedDocument]):
        self.documents = documents
        self.cleaner = DocumentProcessor()
    
    def generate(self) -> str:
        """Generate complete markdown document"""
        md_content = []
        
        # Header
        md_content.append(self._generate_header())
        
        # Table of Contents
        md_content.append(self._generate_toc())
        
        # Process documents by category
        md_content.append("\n## Dokumen yang Diproses\n")
        
        for doc in self.documents:
            md_content.append(self._generate_document_section(doc))
        
        # Summary
        md_content.append(self._generate_summary())
        
        # Footer
        md_content.append(self._generate_footer())
        
        return "\n".join(md_content)
    
    def _generate_header(self) -> str:
        """Generate document header"""
        return f"""# Dokumen SMAN 1 Baleendah - Terproses untuk RAG AI

**Tanggal Pemrosesan:** {datetime.now().strftime('%d %B %Y')}
**Total Dokumen:** {len(self.documents)}
**Versi Prosesor:** 1.0

> Dokumen ini digunakan sebagai knowledge base untuk AI chatbot dan sistem问答 sistem SMAN 1 Baleendah.
"""
    
    def _generate_toc(self) -> str:
        """Generate table of contents"""
        categories = set(doc.category for doc in self.documents)
        toc_lines = ["## Daftar Isi Kategori\n"]
        
        for category in sorted(categories):
            toc_lines.append(f"- **{category}**")
            for doc in [d for d in self.documents if d.category == category]:
                toc_lines.append(f"  - {doc.filename}")
        
        return "\n".join(toc_lines)
    
    def _generate_document_section(self, doc: ProcessedDocument) -> str:
        """Generate section for single document"""
        # Divider
        section = []
        section.append("---\n")
        
        # Header with category tag
        section.append(f"## 📄 {doc.filename}")
        section.append(f"**Kategori:** {doc.category}")
        section.append(f"**Proses:** {doc.processing_date}")
        
        # Metadata
        if doc.metadata:
            section.append(f"\n**Metadata:**")
            for key, value in doc.metadata.items():
                section.append(f"  - {key}: {value}")
        
        # Image references (if any)
        if doc.images:
            section.append(f"\n**Gambar Terkait:**")
            for img_path in doc.images:
                section.append(f"  - ![{doc.filename}]({img_path})")
        
        # Content
        section.append(f"\n### Konten Dokumen\n")
        
        # Clean and format content
        clean_content = self.cleaner.clean_text(doc.content)
        
        # Limit content for readability
        if len(clean_content) > 5000:
            section.append(clean_content[:5000])
            section.append("\n\n*... (konten dipotong, lihat file asli untuk konten lengkap)*")
        else:
            section.append(clean_content)
        
        return "\n".join(section)
    
    def _generate_summary(self) -> str:
        """Generate summary statistics"""
        total_chars = sum(len(doc.content) for doc in self.documents)
        categories = {}
        for doc in self.documents:
            categories[doc.category] = categories.get(doc.category, 0) + 1
        
        summary_lines = [
            "## Ringkasan Proses",
            "",
            f"- **Total Dokumen:** {len(self.documents)}",
            f"- **Total Karakter:** {total_chars:,}",
            f"- **Kategori:** {len(categories)}",
            "",
            "### Distribusi Kategori"
        ]
        
        for category, count in sorted(categories.items()):
            summary_lines.append(f"- **{category}:** {count} dokumen")
        
        summary_lines.extend([
            "",
            "### Tipe File",
        ])
        
        file_types = {}
        for doc in self.documents:
            suffix = Path(doc.filename).suffix.upper() or 'unknown'
            file_types[suffix] = file_types.get(suffix, 0) + 1
        
        for ftype, count in sorted(file_types.items()):
            summary_lines.append(f"- **{ftype}:** {count} file")
        
        summary_lines.extend([
            "",
            f"- **Total Gambar:** {sum(len(doc.images) for doc in self.documents)}",
        ])
        
        return "\n".join(summary_lines)
    
    def _generate_footer(self) -> str:
        """Generate document footer"""
        return f"""

---

**Dokumen ini diproses secara otomatis oleh sistem SMAN 1 Baleendah  
Untuk informasi lebih lanjut, hubungi admin sekolah.**

Note: Dokumen ini untuk keperluan internal sistem RAG AI dan tidak untuk distribusi publik tanpa izin.
"""


def main():
    """Main execution function"""
    print("=" * 60)
    print("  Dokumen SMAN 1 BALEENDAH - Processor")
    print("=" * 60)
    print()
    
    try:
        # Initialize processor
        processor = DocumentProcessor()
        
        # Process all documents
        documents = processor.process_directory(DOC_ROOT)
        
        if not documents:
            print("Tidak ada dokumen yang dapat diproses.")
            return
        
        # Generate markdown
        print("Generating markdown...")
        generator = MarkdownGenerator(documents)
        markdown_content = generator.generate()
        
        # Write to file
        OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        print(f"Markdown successfully generated!")
        print(f"Output file: {OUTPUT_FILE}")
        print(f"File size: {OUTPUT_FILE.stat().st_size / 1024:.1f} KB")
        print()
        print("=" * 60)
        print("  PROSES SELESAI!")
        print("=" * 60)
        
        # Print quick summary
        print("Ringkasan:")
        print(f"   - Total dokumen: {len(documents)}")
        print(f"   - Categories: {set(doc.category for doc in documents)}")
        print(f"   - Images: {sum(len(doc.images) for doc in documents)}")
        
    except Exception as e:
        print(f"Error during processing: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
